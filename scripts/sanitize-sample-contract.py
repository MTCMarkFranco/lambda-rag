"""Sanitize the airline-customer test contract into a generic Contoso sample.

Reads the synthetic test contract from the customer engagement repo, scrubs
airline / customer-specific language, and writes a generic .docx into
samples/contracts/ for use as the bundled markup-mode demo input.
"""
from pathlib import Path
from docx import Document

SRC = Path(r"C:\Users\marfra\source\repos\air-canada\test-data\contract_original.docx")
DEST = Path(r"C:\Users\marfra\source\repos\lambda-rag\samples\contracts\contoso-sample-contract.docx")

REPLACEMENTS = [
    ("Air Canada", "Contoso"),
    ("AIR CANADA", "CONTOSO"),
    ("passenger booking patterns and operational data",
     "customer transaction patterns and operational data"),
    ("routing and scheduling optimization", "operational optimization"),
    ("January 15, 2026", "[Effective Date]"),
]


def replace_in_paragraph(paragraph, find, replace):
    if find not in paragraph.text:
        return
    if not paragraph.runs:
        return
    new_text = paragraph.text.replace(find, replace)
    first = paragraph.runs[0]
    first.text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def main():
    doc = Document(SRC)
    for p in doc.paragraphs:
        for find, replace in REPLACEMENTS:
            replace_in_paragraph(p, find, replace)
    DEST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DEST)
    print(f"Wrote {DEST} ({DEST.stat().st_size} bytes)")
    print("\n--- sanitized text ---")
    out = Document(DEST)
    for i, p in enumerate(out.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")


if __name__ == "__main__":
    main()
