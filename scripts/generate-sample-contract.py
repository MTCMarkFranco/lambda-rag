"""Generate the bundled sample contract DOCX from contract.md.

Fully synthetic input (Acme Corp. / Vendor Inc., Delaware governing law).
No customer-derived content. Run with:

    python scripts/generate-sample-contract.py
"""
from pathlib import Path
import re

from docx import Document

REPO = Path(__file__).resolve().parents[1]
SRC = REPO / "samples" / "contracts" / "contract.md"
DEST = REPO / "samples" / "contracts" / "contoso-sample-contract.docx"


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    doc = Document()
    for line in text.splitlines():
        if not line.strip():
            doc.add_paragraph("")
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            continue
        para = doc.add_paragraph()
        # split on **bold** spans
        for i, chunk in enumerate(re.split(r"\*\*([^*]+)\*\*", line)):
            run = para.add_run(chunk)
            if i % 2 == 1:
                run.bold = True
    doc.save(DEST)
    print(f"Wrote {DEST.relative_to(REPO)} ({DEST.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
